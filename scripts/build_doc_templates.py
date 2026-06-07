"""
Создание организационных Word-шаблонов с плейсхолдерами {{KEY}}.

Запуск из корня:
  python scripts/build_doc_templates.py          # только если файла нет
  python scripts/build_doc_templates.py --rebuild  # перезаписать базовые шаблоны
  python scripts/build_doc_templates.py --list   # список плейсхолдеров для фирменной вёрстки
"""
from __future__ import annotations

import argparse
import os
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "server"))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

from services.docx_templates import ensure_parent_dir, write_minimal_template  # noqa: E402

ORG_LINE = "ООН ПКР — Отраслевые общественные награды промышленного комплекса России"

TEMPLATES = [
    (
        "data/templates/Удостоверение к награде.docx",
        "УДОСТОВЕРЕНИЕ К НАГРАДЕ",
        ["FULL_NAME", "AWARD_NAME", "CERT_NUMBER", "PROTOCOL_NUMBER", "REG_DATE", "SIGNER"],
    ),
    (
        "data/templates/Бюллетень голосования.docx",
        "БЮЛЛЕТЕНЬ ГОЛОСОВАНИЯ",
        ["BULLETIN_NUMBER", "BULLETIN_TYPE", "VOTING_START", "VOTING_END", "ADDRESS", "QUESTIONS"],
    ),
    (
        "data/templates/Протокол краткий.docx",
        "ПРОТОКОЛ (КРАТКИЙ)",
        ["PROTOCOL_NUMBER", "PROTOCOL_DATE", "BULLETIN_NUMBER", "STATUS", "RESULTS_SUMMARY"],
    ),
    (
        "data/templates/Протокол подробный.docx",
        "ПРОТОКОЛ (ПОДРОБНЫЙ)",
        ["PROTOCOL_NUMBER", "PROTOCOL_DATE", "BULLETIN_NUMBER", "STATUS", "DETAILS", "RESULTS_TABLE"],
    ),
    (
        "data/templates/Выписка из протокола — медаль.docx",
        "ВЫПИСКА ИЗ ПРОТОКОЛА (МЕДАЛЬ / ЗНАК)",
        ["EXTRACT_NUMBER", "PROTOCOL_NUMBER", "PROTOCOL_DATE", "FULL_NAME", "AWARD_NAME", "EXTRACT_DATE", "DETAILS"],
    ),
    (
        "data/templates/Выписка из протокола — ППЗ.docx",
        "ВЫПИСКА ИЗ ПРОТОКОЛА (ППЗ)",
        ["EXTRACT_NUMBER", "PROTOCOL_NUMBER", "PROTOCOL_DATE", "FULL_NAME", "AWARD_NAME", "EXTRACT_DATE", "DETAILS"],
    ),
    (
        "data/templates/Представление ППЗ.docx",
        "ПРЕДСТАВЛЕНИЕ НА НАГРАЖДЕНИЕ (ППЗ)",
        ["SUBMISSION_NUMBER", "DATE", "AUTHORIZED", "FULL_NAME", "AWARD_NAME", "DETAILS"],
    ),
    (
        "data/templates/Мониторинг ответов.docx",
        "МОНИТОРИНГ ОТВЕТОВ НК",
        ["BULLETIN_NUMBER", "REQUIRED", "RECEIVED", "MEMBERS_TABLE"],
    ),
]

FIELD_LABELS = {
    "FULL_NAME": "ФИО",
    "AWARD_NAME": "Награда",
    "CERT_NUMBER": "№ удостоверения",
    "PROTOCOL_NUMBER": "№ протокола",
    "REG_DATE": "Дата регистрации",
    "SIGNER": "Подписант",
    "BULLETIN_NUMBER": "№ бюллетеня",
    "BULLETIN_TYPE": "Тип",
    "VOTING_START": "Начало голосования",
    "VOTING_END": "Окончание голосования",
    "ADDRESS": "Адрес для ответа",
    "QUESTIONS": "Вопросы",
    "PROTOCOL_DATE": "Дата протокола",
    "STATUS": "Статус",
    "RESULTS_SUMMARY": "Итоги (кратко)",
    "DETAILS": "Примечание",
    "RESULTS_TABLE": "Таблица результатов",
    "EXTRACT_NUMBER": "№ выписки",
    "EXTRACT_DATE": "Дата выписки",
    "SUBMISSION_NUMBER": "№ представления",
    "DATE": "Дата",
    "AUTHORIZED": "Уполномоченный НК",
    "REQUIRED": "Требуется ответов",
    "RECEIVED": "Получено ответов",
    "MEMBERS_TABLE": "Таблица членов НК",
}


def _org_header(doc: Document) -> None:
    p = doc.add_paragraph(ORG_LINE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True
    doc.add_paragraph("")


def _title(doc: Document, title: str) -> None:
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _field_block(doc: Document, placeholders: list[str]) -> None:
    for key in placeholders:
        label = FIELD_LABELS.get(key, key)
        doc.add_paragraph(f"{label}: {{{{{key}}}}}")


def build_rich_template(path: str, title: str, placeholders: list[str]) -> None:
    ensure_parent_dir(path)
    doc = Document()
    _org_header(doc)
    _title(doc, title)
    if title.startswith("БЮЛЛЕТЕНЬ"):
        doc.add_paragraph("№ {{BULLETIN_NUMBER}}    Тип: {{BULLETIN_TYPE}}")
        doc.add_paragraph("Период: {{VOTING_START}} — {{VOTING_END}}")
        doc.add_paragraph("Адрес: {{ADDRESS}}")
        doc.add_paragraph("")
        doc.add_paragraph("Вопросы:")
        doc.add_paragraph("{{QUESTIONS}}")
    elif "МОНИТОРИНГ" in title:
        doc.add_paragraph("Бюллетень № {{BULLETIN_NUMBER}}")
        doc.add_paragraph("Требуется: {{REQUIRED}}    Получено: {{RECEIVED}}")
        doc.add_paragraph("")
        doc.add_paragraph("{{MEMBERS_TABLE}}")
    elif "ПРОТОКОЛ (ПОДРОБНЫЙ)" in title:
        doc.add_paragraph("№ {{PROTOCOL_NUMBER}} от {{PROTOCOL_DATE}}")
        doc.add_paragraph("Бюллетень № {{BULLETIN_NUMBER}}    Статус: {{STATUS}}")
        doc.add_paragraph("{{DETAILS}}")
        doc.add_paragraph("")
        doc.add_paragraph("Результаты:")
        doc.add_paragraph("{{RESULTS_TABLE}}")
    else:
        _field_block(doc, placeholders)
    doc.save(path)


def list_placeholders() -> None:
    print("Плейсхолдеры для фирменных шаблонов (подставьте в свой .docx как {{KEY}}):\n")
    for rel, title, placeholders in TEMPLATES:
        print(f"  {rel}")
        print(f"    {title}")
        for ph in placeholders:
            print(f"      {{{{{ph}}}}} — {FIELD_LABELS.get(ph, ph)}")
        print()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--rebuild", action="store_true", help="Перезаписать шаблоны улучшенной вёрсткой")
    parser.add_argument("--list", action="store_true", help="Список плейсхолдеров")
    args = parser.parse_args()

    os.chdir(ROOT)

    if args.list:
        list_placeholders()
        return 0

    for rel, title, placeholders in TEMPLATES:
        path = os.path.join(ROOT, rel)
        if args.rebuild:
            build_rich_template(path, title, placeholders)
            print(f"REBUILD: {rel}")
        else:
            write_minimal_template(path, title, placeholders)
            print(f"OK: {rel}")
    print("Готово.")
    if not args.rebuild:
        print("Подсказка: python scripts/build_doc_templates.py --rebuild — базовая шапка ООН ПКР")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
