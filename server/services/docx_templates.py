"""Загрузка организационных Word-шаблонов и подстановка плейсхолдеров {{KEY}}."""

import os
import re
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from fastapi import HTTPException

_PLACEHOLDER = re.compile(r"\{\{(\w+)\}\}")

_PRETTY_LABELS = {
    "BULLETIN_NUMBER": "Номер бюллетеня",
    "BULLETIN_TYPE": "Тип бюллетеня",
    "VOTING_START": "Начало голосования",
    "VOTING_END": "Окончание голосования",
    "ADDRESS": "Адрес",
    "QUESTIONS": "Вопросы",
    "PROTOCOL_NUMBER": "Номер протокола",
    "PROTOCOL_DATE": "Дата протокола",
    "STATUS": "Статус",
    "DETAILS": "Комментарий",
    "RESULTS_SUMMARY": "Итоги голосования",
    "RESULTS_TABLE": "Таблица результатов",
    "EXTRACT_NUMBER": "Номер выписки",
    "EXTRACT_DATE": "Дата выписки",
    "FULL_NAME": "ФИО",
    "AWARD_NAME": "Награда",
    "SUBMISSION_NUMBER": "Номер представления",
    "DATE": "Дата",
    "AUTHORIZED": "Уполномоченный представитель",
}


def replace_placeholders(doc: Document, mapping: dict[str, str]) -> int:
    replaced = 0

    def apply_paragraphs(paragraphs):
        nonlocal replaced
        for p in paragraphs:
            if "{{" not in p.text:
                continue
            full_text = "".join(run.text for run in p.runs)
            new_text = full_text
            for m in _PLACEHOLDER.finditer(full_text):
                key = m.group(1)
                new_text = new_text.replace(m.group(0), mapping.get(key, "—"))
                replaced += 1
            if new_text == full_text:
                continue
            if p.runs:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""

    apply_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_paragraphs(cell.paragraphs)
    return replaced


def _open_document(path: str) -> Document:
    """Open .docx and Word template files saved with .docx extension."""
    try:
        return Document(path)
    except ValueError as exc:
        if "wordprocessingml.template.main+xml" not in str(exc):
            raise

    with ZipFile(path) as src:
        buf = BytesIO()
        with ZipFile(buf, "w", ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = src.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    )
                dst.writestr(item, data)
    buf.seek(0)
    return Document(buf)


def _append_fallback_data(doc: Document, label: str, mapping: dict[str, str]) -> None:
    doc.add_page_break()
    try:
        doc.add_heading(f"Данные для документа: {label}", level=2)
    except KeyError:
        p = doc.add_paragraph(f"Данные для документа: {label}")
        if p.runs:
            p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.rows[0].cells[0].text = "Поле"
    table.rows[0].cells[1].text = "Значение"
    for key, value in mapping.items():
        if value in (None, "", "—"):
            continue
        cells = table.add_row().cells
        cells[0].text = _PRETTY_LABELS.get(key, key)
        cells[1].text = str(value)


def load_template_or_raise(path: str, label: str) -> Document:
    if not path or not os.path.isfile(path):
        raise HTTPException(
            status_code=500,
            detail=f"Шаблон «{label}» не найден: {path}. Запустите: python scripts/build_doc_templates.py",
        )
    return _open_document(path)


def render_template(path: str, label: str, mapping: dict[str, str]) -> bytes:
    doc = load_template_or_raise(path, label)
    replaced = replace_placeholders(doc, mapping)
    if replaced == 0:
        _append_fallback_data(doc, label, mapping)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def ensure_parent_dir(path: str) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)


def write_minimal_template(path: str, title: str, placeholders: list[str]) -> None:
    """Создать простой шаблон с плейсхолдерами (если файла ещё нет)."""
    if os.path.isfile(path):
        return
    ensure_parent_dir(path)
    doc = Document()
    doc.add_heading(title, level=1)
    for ph in placeholders:
        doc.add_paragraph(f"{{{{{ph}}}}}")
    doc.save(path)
