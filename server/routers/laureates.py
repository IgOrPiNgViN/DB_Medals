import os
import re
from io import BytesIO

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, Query, status, UploadFile, File
from fastapi.responses import Response
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func
from typing import List, Optional
from datetime import date

from database import get_db
from config import CONSENT_TEMPLATE_PATH
from models.laureate import Laureate, LaureateAward, LaureateLifecycle, LaureateConsentFile
from models.award import Award
from schemas.laureate import (
    LaureateCreate, LaureateUpdate, LaureateRead,
    LaureateAwardCreate, LaureateAwardRead,
    LaureateLifecycleCreate, LaureateLifecycleUpdate, LaureateLifecycleRead,
)

router = APIRouter()


def _get_laureate_or_404(db: Session, laureate_id: int) -> Laureate:
    obj = db.query(Laureate).filter(Laureate.id == laureate_id).first()
    if not obj:
        raise HTTPException(status_code=404, detail="Laureate not found")
    return obj


def _get_laureate_award_or_404(db: Session, la_id: int) -> LaureateAward:
    obj = db.query(LaureateAward).filter(LaureateAward.id == la_id).first()
    if not obj:
        raise HTTPException(status_code=404, detail="LaureateAward not found")
    return obj


# ── Laureate CRUD ───────────────────────────────────────────────────────────

@router.get("/", response_model=List[LaureateRead])
def list_laureates(category: Optional[str] = None, db: Session = Depends(get_db)):
    q = db.query(Laureate)
    if category:
        q = q.filter(Laureate.category == category)
    return q.all()


@router.get("/laureate-awards/by-bulletin")
def list_laureate_awards_by_bulletin_number(
    bulletin_number: str = Query(..., min_length=1, description="Номер бюллетеня (как в карточке связки)"),
    db: Session = Depends(get_db),
):
    """
    Все связки лауреат–награда с указанным номером бюллетеня — для конструктора бюллетеня (раздел «Награждение»).
    """
    bn = (bulletin_number or "").strip()
    if not bn:
        return []
    rows = (
        db.query(LaureateAward)
        .options(
            joinedload(LaureateAward.laureate),
            joinedload(LaureateAward.award),
        )
        .filter(LaureateAward.bulletin_number == bn)
        .order_by(LaureateAward.id)
        .all()
    )
    return [
        {
            "laureate_award_id": la.id,
            "laureate_id": la.laureate_id,
            "full_name": la.laureate.full_name if la.laureate else "",
            "award_name": la.award.name if la.award else "",
            "award_id": la.award_id,
            "bulletin_number": la.bulletin_number,
            "initiator": la.initiator,
        }
        for la in rows
    ]


@router.get("/links/{laureate_award_id}")
def get_laureate_award_context(laureate_award_id: int, db: Session = Depends(get_db)):
    """Контекст связки лауреат–награда (ФИО, название награды) для печати и форм."""
    la = (
        db.query(LaureateAward)
        .options(
            joinedload(LaureateAward.laureate),
            joinedload(LaureateAward.award),
        )
        .filter(LaureateAward.id == laureate_award_id)
        .first()
    )
    if not la:
        raise HTTPException(status_code=404, detail="LaureateAward not found")
    return {
        "laureate_award_id": la.id,
        "laureate_id": la.laureate_id,
        "full_name": la.laureate.full_name if la.laureate else "",
        "award_id": la.award_id,
        "award_name": la.award.name if la.award else "",
        "award_type": la.award.award_type.value if la.award and la.award.award_type else None,
        "assigned_date": la.assigned_date,
        "status": la.status,
        "bulletin_number": la.bulletin_number,
        "initiator": la.initiator,
    }


@router.post("/", response_model=LaureateRead, status_code=status.HTTP_201_CREATED)
def create_laureate(payload: LaureateCreate, db: Session = Depends(get_db)):
    obj = Laureate(**payload.model_dump())
    db.add(obj)
    db.commit()
    db.refresh(obj)
    return obj


@router.get("/reports/awards-laureates")
def awards_laureates_report(db: Session = Depends(get_db)):
    """Все лауреаты, сгруппированные по награде."""
    awards = db.query(Award).options(
        joinedload(Award.laureate_awards).joinedload(LaureateAward.laureate),
    ).all()
    result = []
    for a in awards:
        laureates = []
        for la in a.laureate_awards:
            laureates.append({
                "laureate_award_id": la.id,
                "laureate_id": la.laureate.id,
                "full_name": la.laureate.full_name,
                "category": la.laureate.category.value if la.laureate.category else None,
                "assigned_date": la.assigned_date,
                "status": la.status,
            })
        result.append({
            "award_id": a.id,
            "award_name": a.name,
            "award_type": a.award_type.value if a.award_type else None,
            "laureates": laureates,
        })
    return result


@router.get("/reports/incomplete-lifecycle")
def incomplete_lifecycle_report(db: Session = Depends(get_db)):
    """Лауреаты с незавершённым жизненным циклом."""
    la_list = (
        db.query(LaureateAward)
        .options(
            joinedload(LaureateAward.laureate),
            joinedload(LaureateAward.award),
            joinedload(LaureateAward.lifecycle),
        )
        .all()
    )
    result = []
    for la in la_list:
        lc = la.lifecycle
        if lc is None:
            result.append({
                "laureate_award_id": la.id,
                "laureate_name": la.laureate.full_name,
                "award_name": la.award.name,
                "reason": "lifecycle not created",
            })
            continue
        incomplete_stages = []
        if not lc.nomination_done:
            incomplete_stages.append("nomination")
        if not lc.voting_done:
            incomplete_stages.append("voting")
        if not lc.decision_done:
            incomplete_stages.append("decision")
        if not lc.registration_done:
            incomplete_stages.append("registration")
        if not lc.ceremony_done:
            incomplete_stages.append("ceremony")
        if not lc.publication_done:
            incomplete_stages.append("publication")
        if incomplete_stages:
            result.append({
                "laureate_award_id": la.id,
                "laureate_name": la.laureate.full_name,
                "award_name": la.award.name,
                "incomplete_stages": incomplete_stages,
            })
    return result


@router.get("/reports/statistics")
def statistics_report(
    from_date: Optional[date] = None,
    to_date: Optional[date] = None,
    db: Session = Depends(get_db),
):
    """Статистика по категориям с фильтром по дате."""
    q = db.query(
        Laureate.category,
        func.count(Laureate.id).label("count"),
    )
    if from_date:
        q = q.filter(Laureate.created_at >= from_date)
    if to_date:
        q = q.filter(Laureate.created_at <= to_date)
    rows = q.group_by(Laureate.category).all()
    return [
        {"category": r.category.value if r.category else None, "count": r.count}
        for r in rows
    ]


@router.get("/{laureate_id}", response_model=LaureateRead)
def get_laureate(laureate_id: int, db: Session = Depends(get_db)):
    return _get_laureate_or_404(db, laureate_id)


@router.put("/{laureate_id}", response_model=LaureateRead)
def update_laureate(
    laureate_id: int, payload: LaureateUpdate, db: Session = Depends(get_db),
):
    obj = _get_laureate_or_404(db, laureate_id)
    for key, value in payload.model_dump(exclude_unset=True).items():
        setattr(obj, key, value)
    db.commit()
    db.refresh(obj)
    return obj


@router.delete("/{laureate_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_laureate(laureate_id: int, db: Session = Depends(get_db)):
    obj = _get_laureate_or_404(db, laureate_id)
    db.delete(obj)
    db.commit()


# ── LaureateAward ───────────────────────────────────────────────────────────

@router.post(
    "/{laureate_id}/awards",
    response_model=LaureateAwardRead,
    status_code=status.HTTP_201_CREATED,
)
def link_award(
    laureate_id: int, payload: LaureateAwardCreate, db: Session = Depends(get_db),
):
    _get_laureate_or_404(db, laureate_id)
    obj = LaureateAward(**payload.model_dump())
    obj.laureate_id = laureate_id
    db.add(obj)
    db.flush()
    if obj.bulletin_number:
        lc = (
            db.query(LaureateLifecycle)
            .filter(LaureateLifecycle.laureate_award_id == obj.id)
            .first()
        )
        if lc is not None:
            lc.voting_bulletin_number = obj.bulletin_number
    db.commit()
    db.refresh(obj)
    return obj


@router.get("/{laureate_id}/awards", response_model=List[LaureateAwardRead])
def list_laureate_awards(laureate_id: int, db: Session = Depends(get_db)):
    _get_laureate_or_404(db, laureate_id)
    return db.query(LaureateAward).filter(
        LaureateAward.laureate_id == laureate_id,
    ).all()


# ── Lifecycle ───────────────────────────────────────────────────────────────

@router.post(
    "/{laureate_award_id}/lifecycle",
    response_model=LaureateLifecycleRead,
    status_code=status.HTTP_201_CREATED,
)
def create_lifecycle(
    laureate_award_id: int,
    payload: LaureateLifecycleCreate,
    db: Session = Depends(get_db),
):
    la = _get_laureate_award_or_404(db, laureate_award_id)
    existing = db.query(LaureateLifecycle).filter(
        LaureateLifecycle.laureate_award_id == laureate_award_id,
    ).first()
    if existing:
        raise HTTPException(status_code=409, detail="Lifecycle already exists")
    obj = LaureateLifecycle(**payload.model_dump())
    obj.laureate_award_id = laureate_award_id
    if not obj.voting_bulletin_number and la.bulletin_number:
        obj.voting_bulletin_number = la.bulletin_number
    db.add(obj)
    db.commit()
    db.refresh(obj)
    return obj


@router.get("/{laureate_award_id}/lifecycle", response_model=LaureateLifecycleRead)
def get_lifecycle(laureate_award_id: int, db: Session = Depends(get_db)):
    obj = db.query(LaureateLifecycle).filter(
        LaureateLifecycle.laureate_award_id == laureate_award_id,
    ).first()
    if not obj:
        raise HTTPException(status_code=404, detail="Lifecycle not found")
    return obj


@router.put("/{laureate_award_id}/lifecycle", response_model=LaureateLifecycleRead)
def update_lifecycle(
    laureate_award_id: int,
    payload: LaureateLifecycleUpdate,
    db: Session = Depends(get_db),
):
    obj = db.query(LaureateLifecycle).filter(
        LaureateLifecycle.laureate_award_id == laureate_award_id,
    ).first()
    if not obj:
        raise HTTPException(status_code=404, detail="Lifecycle not found")
    for key, value in payload.model_dump(exclude_unset=True).items():
        setattr(obj, key, value)
    db.commit()
    db.refresh(obj)
    return obj


# ── Consent PD (персональные данные) ─────────────────────────────────────────


@router.get("/{laureate_award_id}/consent/file/info")
def get_consent_file_info(laureate_award_id: int, db: Session = Depends(get_db)):
    obj = db.query(LaureateConsentFile).filter(
        LaureateConsentFile.laureate_award_id == laureate_award_id,
    ).first()
    if not obj:
        return {"exists": False}
    return {
        "exists": True,
        "filename": obj.filename,
        "content_type": obj.content_type,
        "size": len(obj.data or b""),
        "uploaded_at": obj.uploaded_at.isoformat() if obj.uploaded_at else None,
    }


@router.get("/{laureate_award_id}/consent/file")
def download_consent_file(laureate_award_id: int, db: Session = Depends(get_db)):
    obj = db.query(LaureateConsentFile).filter(
        LaureateConsentFile.laureate_award_id == laureate_award_id,
    ).first()
    if not obj:
        raise HTTPException(status_code=404, detail="Consent file not found")
    headers = {"Content-Disposition": f'attachment; filename="{obj.filename}"'}
    return Response(
        content=obj.data,
        media_type=obj.content_type or "application/octet-stream",
        headers=headers,
    )


@router.post("/{laureate_award_id}/consent/file", status_code=status.HTTP_204_NO_CONTENT)
async def upload_consent_file(
    laureate_award_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    _get_laureate_award_or_404(db, laureate_award_id)
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    obj = db.query(LaureateConsentFile).filter(
        LaureateConsentFile.laureate_award_id == laureate_award_id,
    ).first()
    if obj:
        obj.filename = file.filename or "consent.bin"
        obj.content_type = file.content_type
        obj.data = data
    else:
        obj = LaureateConsentFile(
            laureate_award_id=laureate_award_id,
            filename=file.filename or "consent.bin",
            content_type=file.content_type,
            data=data,
        )
        db.add(obj)
    db.commit()
    return None


@router.delete("/{laureate_award_id}/consent/file", status_code=status.HTTP_204_NO_CONTENT)
def delete_consent_file(laureate_award_id: int, db: Session = Depends(get_db)):
    obj = db.query(LaureateConsentFile).filter(
        LaureateConsentFile.laureate_award_id == laureate_award_id,
    ).first()
    if obj:
        db.delete(obj)
        db.commit()
    return None


def _fill_doc_placeholders(doc: Document, full_name: str) -> None:
    # В шаблоне стоят линии из подчёркиваний — заменяем их на ФИО.
    pattern = re.compile(r"_{3,}")

    def replace_in_paragraphs(paragraphs):
        for p in paragraphs:
            text = p.text
            if not text or "_" not in text:
                continue
            p.text = pattern.sub(full_name, text)

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)


@router.get("/{laureate_award_id}/consent/generate")
def generate_consent_docx(laureate_award_id: int, db: Session = Depends(get_db)):
    la = _get_laureate_award_or_404(db, laureate_award_id)
    full_name = la.laureate.full_name if la.laureate else ""
    if not full_name:
        raise HTTPException(status_code=400, detail="Laureate full_name is empty")

    template_path = CONSENT_TEMPLATE_PATH
    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail=f"Consent template not found: {template_path}")

    doc = Document(template_path)
    _fill_doc_placeholders(doc, full_name)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r"[\\\\/:*?\"<>|]+", "_", full_name).strip() or "laureate"
    filename = f"Согласие ПД — {safe_name}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
